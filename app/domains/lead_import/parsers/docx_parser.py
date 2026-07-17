"""DOCX parser using python-docx."""

from pathlib import Path
from docx import Document as DocxDocument
from lxml import etree
from app.domains.lead_import.parsers.base import BaseParser, ExtractedDocument


class DOCXParser(BaseParser):
    """Parse .docx files — extracts paragraph text and table content in order."""

    async def parse(self, file_path: str) -> ExtractedDocument:
        path = Path(file_path)
        doc = DocxDocument(file_path)
        result = ExtractedDocument(metadata={"filename": path.name})
        text_parts = []
        tables = []

        # Walk XML body in document order to interleave paragraphs and tables
        body = doc.element.body
        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        for child in body:
            tag = etree.QName(child).localname
            if tag == "p":
                texts = child.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
                line = "".join(t.text or "" for t in texts).strip()
                if line:
                    text_parts.append(line)
            elif tag == "tbl":
                rows = child.findall(".//w:tr", nsmap)
                table_data = []
                for row in rows:
                    cells = row.findall(".//w:tc//w:t", nsmap)
                    row_data = [c.text or "" for c in cells]
                    if any(c.strip() for c in row_data):
                        table_data.append(row_data)
                if table_data:
                    tables.append(table_data)
                    for data_row in table_data:
                        text_parts.append(" | ".join(data_row))

        result.text = "\n".join(text_parts)
        result.tables = tables
        result.pages = 1
        result.metadata["paragraphs"] = sum(1 for c in body if etree.QName(c).localname == "p")
        result.metadata["tables"] = len(tables)

        return result
